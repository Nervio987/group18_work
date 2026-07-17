"""
文档生成模块：将 Markdown 文本转换为 Word (.docx) 或 PDF 文档
"""
import os
import re
import uuid
import time
import threading
from pathlib import Path
from typing import Tuple

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF

# 临时文件目录
DOCS_DIR = Path(os.getenv("DOCS_OUTPUT_DIR", "./generated_docs"))
DOCS_DIR.mkdir(parents=True, exist_ok=True)

# 文件清理配置
FILE_TTL_SECONDS = 3600  # 文件保留 1 小时
MAX_FILES = 200  # 最多保留 200 个文件
_file_registry: dict[str, dict] = {}
_registry_lock = threading.Lock()


def _register_file(file_path: Path) -> str:
    """注册生成的文件，返回 file_id"""
    file_id = uuid.uuid4().hex
    with _registry_lock:
        _file_registry[file_id] = {
            "path": str(file_path),
            "filename": file_path.name,
            "created_at": time.time(),
        }
        # 清理过期和超额文件
        _cleanup_expired()
    return file_id


def _cleanup_expired():
    """清理过期的生成文件"""
    now = time.time()
    # 删除过期文件
    expired = [fid for fid, info in _file_registry.items()
               if now - info["created_at"] > FILE_TTL_SECONDS]
    for fid in expired:
        info = _file_registry.pop(fid, None)
        if info:
            try:
                os.remove(info["path"])
            except OSError:
                pass

    # 如果仍超过上限，删除最旧的文件
    if len(_file_registry) > MAX_FILES:
        sorted_files = sorted(_file_registry.items(),
                             key=lambda x: x[1]["created_at"])
        to_delete = len(_file_registry) - MAX_FILES
        for fid, info in sorted_files[:to_delete]:
            _file_registry.pop(fid, None)
            try:
                os.remove(info["path"])
            except OSError:
                pass


def get_file_path(file_id: str) -> str | None:
    """根据 file_id 获取文件路径"""
    with _registry_lock:
        info = _file_registry.get(file_id)
        if not info:
            return None
        # 检查是否过期
        if time.time() - info["created_at"] > FILE_TTL_SECONDS:
            _file_registry.pop(file_id, None)
            try:
                os.remove(info["path"])
            except OSError:
                pass
            return None
        return info["path"]


# ============ Word 文档生成 ============

def _add_inline_runs(paragraph, text: str, base_size: int = 11):
    """处理 Markdown 内联格式（粗体、斜体、代码、链接）"""
    # 解析顺序: 链接 -> 粗体 -> 斜体 -> 行内代码
    # 使用更稳健的逐段扫描方法
    patterns = [
        (re.compile(r"\[([^\]]+)\]\(([^)]+)\)"), "link"),
        (re.compile(r"\*\*([^*]+)\*\*"), "bold"),
        (re.compile(r"\*([^*]+)\*"), "italic"),
        (re.compile(r"`([^`]+)`"), "code"),
    ]
    # 简化处理：使用占位符替换法
    placeholders = []

    def make_placeholder(style: str, content: str, url: str = "") -> str:
        idx = len(placeholders)
        token = f"\x00PH{idx}\x00"
        placeholders.append((style, content, url))
        return token

    tmp = text
    for pat, style in patterns:
        if style == "link":
            tmp = pat.sub(lambda m: make_placeholder("link", m.group(1), m.group(2)), tmp)
        else:
            tmp = pat.sub(lambda m: make_placeholder(style, m.group(1)), tmp)

    # 按占位符切分
    parts = re.split(r"(\x00PH\d+\x00)", tmp)

    for part in parts:
        if not part:
            continue
        m = re.fullmatch(r"\x00PH(\d+)\x00", part)
        if m:
            style, content, url = placeholders[int(m.group(1))]
            run = paragraph.add_run(content)
            run.font.size = Pt(base_size)
            if style == "bold":
                run.bold = True
            elif style == "italic":
                run.italic = True
            elif style == "code":
                run.font.name = "Consolas"
                run.font.size = Pt(base_size - 1)
                run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
            elif style == "link":
                run.font.color.rgb = RGBColor(0x4F, 0x6E, 0xF7)
                run.underline = True
        else:
            run = paragraph.add_run(part)
            run.font.size = Pt(base_size)


def _parse_table_block(lines: list[str], doc: Document):
    """解析 Markdown 表格并添加到 Word 文档"""
    if not lines or "|" not in lines[0]:
        return
    # 跳过分隔行（|---|...）
    header_cells = [c.strip() for c in lines[0].strip().strip("|").split("|")]
    rows = []
    for line in lines[2:]:
        if "|" not in line:
            break
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)

    table = doc.add_table(rows=1 + len(rows), cols=len(header_cells))
    table.style = "Light Grid Accent 1"
    # 表头
    for i, cell_text in enumerate(header_cells):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        _add_inline_runs(p, cell_text, base_size=11)
        for r in p.runs:
            r.bold = True
    # 数据行
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            if c_idx >= len(table.rows[r_idx + 1].cells):
                continue
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            _add_inline_runs(p, cell_text, base_size=10)


def _parse_code_block(lines: list[str], doc: Document):
    """解析代码块并添加到 Word 文档"""
    code_text = "\n".join(lines)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
    # 添加浅灰背景（通过段落边框模拟）
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F3F4F6')
    pPr.append(shd)


def markdown_to_docx(markdown_text: str, title: str = "文档") -> str:
    """将 Markdown 文本转换为 Word 文档，返回文件路径"""
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(11)

    # 标题
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(20)
    title_run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

    # 分隔线
    doc.add_paragraph("_" * 50)

    # 解析 Markdown
    lines = markdown_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 跳过空行
        if not stripped:
            doc.add_paragraph()
            i += 1
            continue

        # 标题
        if stripped.startswith("# "):
            p = doc.add_paragraph()
            r = p.add_run(stripped[2:])
            r.bold = True
            r.font.size = Pt(18)
            i += 1
            continue
        if stripped.startswith("## "):
            p = doc.add_paragraph()
            r = p.add_run(stripped[3:])
            r.bold = True
            r.font.size = Pt(15)
            r.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
            i += 1
            continue
        if stripped.startswith("### "):
            p = doc.add_paragraph()
            r = p.add_run(stripped[4:])
            r.bold = True
            r.font.size = Pt(13)
            r.font.color.rgb = RGBColor(0x37, 0x41, 0x51)
            i += 1
            continue
        if stripped.startswith("#### "):
            p = doc.add_paragraph()
            r = p.add_run(stripped[5:])
            r.bold = True
            r.font.size = Pt(12)
            i += 1
            continue

        # 代码块
        if stripped.startswith("```"):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            _parse_code_block(code_lines, doc)
            i += 1  # 跳过结束 ```
            continue

        # 表格
        if "|" in stripped and i + 1 < len(lines) and re.match(r"^\|?[\s:|-]+\|?$", lines[i + 1].strip()):
            table_lines = []
            while i < len(lines) and "|" in lines[i]:
                table_lines.append(lines[i])
                i += 1
            _parse_table_block(table_lines, doc)
            continue

        # 引用
        if stripped.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            r = p.add_run('│ ' + stripped[2:])
            r.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
            r.italic = True
            i += 1
            continue

        # 无序列表
        if re.match(r"^[\-\*]\s+", stripped):
            text = re.sub(r"^[\-\*]\s+", "", stripped)
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_runs(p, text, base_size=11)
            i += 1
            continue

        # 有序列表
        m = re.match(r"^(\d+)\.\s+(.+)", stripped)
        if m:
            text = m.group(2)
            p = doc.add_paragraph(style="List Number")
            _add_inline_runs(p, text, base_size=11)
            i += 1
            continue

        # 分隔线
        if re.match(r"^[-*_]{3,}$", stripped):
            p = doc.add_paragraph("─" * 60)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        # 普通段落（合并连续非空行）
        para_lines = [stripped]
        i += 1
        while i < len(lines) and lines[i].strip() and not _is_block_start(lines[i].strip()):
            para_lines.append(lines[i].strip())
            i += 1
        p = doc.add_paragraph()
        _add_inline_runs(p, " ".join(para_lines), base_size=11)

    # 保存
    safe_title = re.sub(r'[\\/:*?"<>|]', '_', title)[:50] or "文档"
    file_name = f"{safe_title}_{uuid.uuid4().hex[:8]}.docx"
    file_path = DOCS_DIR / file_name
    doc.save(str(file_path))
    return str(file_path)


def _is_block_start(s: str) -> bool:
    """判断是否为新的 Markdown 块开始"""
    if not s:
        return False
    if s.startswith("#") or s.startswith("```") or s.startswith("> "):
        return True
    if re.match(r"^[\-\*]\s+", s):
        return True
    if re.match(r"^\d+\.\s+", s):
        return True
    if re.match(r"^[-*_]{3,}$", s):
        return True
    if "|" in s and re.match(r"^\|?[\s:|-]+\|?$", s):
        return True
    return False


# ============ PDF 文档生成 ============

# 中文字体路径（按优先级尝试，纯 TTF 优先于 TTC）
CHINESE_FONT_PATHS = [
    r"C:\Windows\Fonts\simhei.ttf",
    r"C:\Windows\Fonts\msyh.ttc",
    r"C:\Windows\Fonts\NotoSansSC-VF.ttf",
    r"C:\Windows\Fonts\Noto Sans SC (TrueType).otf",
    r"C:\Windows\Fonts\simsun.ttc",
    "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
    "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc",
    "/System/Library/Fonts/PingFang.ttc",
]

_chinese_font_path: str | None = None
_chinese_font_loaded = False


def _find_chinese_font() -> str | None:
    """查找可用的中文字体"""
    global _chinese_font_loaded
    if _chinese_font_loaded:
        return _chinese_font_path
    for path in CHINESE_FONT_PATHS:
        if os.path.exists(path):
            _chinese_font_path = path
            break
    _chinese_font_loaded = True
    return _chinese_font_path


def _strip_markdown_for_pdf(text: str) -> str:
    """为 PDF 简化 Markdown 格式（去除复杂标记，保留基本格式）"""
    # 移除代码块标记
    text = re.sub(r"```[a-zA-Z]*\n?", "", text)
    text = re.sub(r"```", "", text)
    # 链接 → 显示文本
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    # 粗体/斜体 → 保留文本
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    # 行内代码 → 保留文本
    text = re.sub(r"`([^`]+)`", r"\1", text)
    # 引用
    text = re.sub(r"^>\s*", "  ", text, flags=re.MULTILINE)
    return text


def markdown_to_pdf(markdown_text: str, title: str = "文档") -> str:
    """将 Markdown 文本转换为 PDF 文档，返回文件路径"""
    try:
        return _markdown_to_pdf_impl(markdown_text, title)
    except Exception as e:
        print(f"[DocGen] PDF 生成失败: {type(e).__name__}: {e}")
        raise


def _markdown_to_pdf_impl(markdown_text: str, title: str = "文档") -> str:
    """PDF 生成实现"""
    pdf = FPDF(orientation="P", unit="mm", format="A4")
    pdf.set_margins(left=15, top=15, right=15)
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    page_width = pdf.w - pdf.l_margin - pdf.r_margin  # 180mm

    # 加载中文字体
    font_path = _find_chinese_font()
    use_chinese = False
    if font_path:
        try:
            pdf.add_font("zh", "", font_path)
            pdf.add_font("zh", "B", font_path)
            pdf.set_font("zh", size=10)
            use_chinese = True
        except Exception as e:
            print(f"[DocGen] 中文字体加载失败: {e}，fallback 到 Helvetica")
            pdf.set_font("Helvetica", size=10)
            page_width = pdf.w - pdf.l_margin - pdf.r_margin
    else:
        print("[DocGen] 未找到中文字体，使用 Helvetica（中文可能显示为方块）")
        pdf.set_font("Helvetica", size=10)

    def _set_font(style="", size=10):
        """统一设置字体（中文优先）"""
        if use_chinese:
            pdf.set_font("zh", style, size)
        else:
            # Helvetica 不支持中文，超长字符可能溢出，做基础防护
            if style == "B":
                pdf.set_font("Helvetica", "B", size)
            else:
                pdf.set_font("Helvetica", size)

    def _safe_text(text: str) -> str:
        """PDF 安全的文本处理：避免 latin-1 编码错误"""
        if use_chinese:
            return text
        # 不支持中文时，移除中文字符避免编码错误
        return re.sub(r"[^\x00-\xff]+", "", text)

    # 标题
    _set_font("B", 18)
    pdf.cell(page_width, 12, _safe_text(_strip_markdown_for_pdf(title)), ln=1, align="C")
    pdf.ln(4)

    # 分隔线
    pdf.line(15, pdf.get_y(), 195, pdf.get_y())
    pdf.ln(6)

    # 解析 Markdown 行
    lines = _strip_markdown_for_pdf(markdown_text).split("\n")
    in_code_block = False

    for line in lines:
        stripped = line.strip()

        # 跳过空行
        if not stripped:
            pdf.ln(3)
            continue

        # 代码块开始/结束
        if stripped.startswith("```"):
            in_code_block = not in_code_block
            if use_chinese:
                pdf.set_font("zh", size=9)
            else:
                pdf.set_font("Courier", size=9)
            pdf.set_text_color(31, 41, 55)
            pdf.set_fill_color(243, 244, 246)
            continue

        if in_code_block:
            # 代码行
            pdf.multi_cell(page_width, 4.5, _safe_text(line), fill=True)
            continue

        # 恢复默认颜色和字体
        pdf.set_text_color(0, 0, 0)
        _set_font("", 10)

        # 标题层级
        if stripped.startswith("#### "):
            _set_font("B", 11)
            pdf.multi_cell(page_width, 6, "  " + _safe_text(stripped[5:]))
            pdf.ln(1)
            continue
        if stripped.startswith("### "):
            _set_font("B", 12)
            pdf.multi_cell(page_width, 6.5, "  " + _safe_text(stripped[4:]))
            pdf.ln(1)
            continue
        if stripped.startswith("## "):
            _set_font("B", 14)
            pdf.multi_cell(page_width, 8, _safe_text(stripped[3:]))
            pdf.ln(2)
            continue
        if stripped.startswith("# "):
            _set_font("B", 16)
            pdf.multi_cell(page_width, 9, _safe_text(stripped[2:]))
            pdf.ln(2)
            continue

        # 引用
        if stripped.startswith(">"):
            pdf.set_text_color(107, 114, 128)
            text = re.sub(r"^>\s*", "", stripped)
            _set_font("", 9)
            pdf.multi_cell(page_width, 5.5, "│ " + _safe_text(text))
            pdf.set_text_color(0, 0, 0)
            continue

        # 无序列表
        m = re.match(r"^[\-\*]\s+(.+)", stripped)
        if m:
            text = m.group(1)
            pdf.multi_cell(page_width, 5.5, "  • " + _safe_text(text))
            continue

        # 有序列表
        m = re.match(r"^(\d+)\.\s+(.+)", stripped)
        if m:
            text = m.group(2)
            pdf.multi_cell(page_width, 5.5, f"  {m.group(1)}. " + _safe_text(text))
            continue

        # 分隔线
        if re.match(r"^[-*_]{3,}$", stripped):
            y = pdf.get_y()
            pdf.line(15, y, 195, y)
            pdf.ln(3)
            continue

        # 普通段落
        pdf.multi_cell(page_width, 5.5, _safe_text(stripped))

    # 保存
    safe_title = re.sub(r'[\\/:*?"<>|]', '_', title)[:50] or "文档"
    file_name = f"{safe_title}_{uuid.uuid4().hex[:8]}.pdf"
    file_path = DOCS_DIR / file_name
    pdf.output(str(file_path))
    return str(file_path)


# ============ 统一入口 ============

def generate_document(content: str, format: str = "word", title: str = "文档") -> Tuple[str, str, str]:
    """
    生成文档的统一入口

    Args:
        content: Markdown 格式的内容
        format: "word" 或 "pdf"
        title: 文档标题

    Returns:
        (file_id, filename, file_path)
    """
    if format == "word":
        file_path = markdown_to_docx(content, title)
    elif format == "pdf":
        file_path = markdown_to_pdf(content, title)
    else:
        raise ValueError(f"不支持的格式: {format}，仅支持 word / pdf")

    file_id = _register_file(Path(file_path))
    return file_id, Path(file_path).name, file_path
